#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_docx.py — Generator Laporan KP Final (.docx)
=====================================================
Script ini menghasilkan file Laporan_KP_Final.docx dengan format sesuai
spesifikasi penulisan akademik (margin, font, spasi, heading, tabel, dll.)
dan konten yang telah disinkronisasi dengan kode proyek aktual.
"""

from docx import Document
from docx.shared import Cm, Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================================
#  CONSTANTS
# ============================================================================
FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5
SINGLE_SPACING = 1.0
FIRST_LINE_INDENT = Cm(1.27)
OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Laporan_KP_Final.docx')


# ============================================================================
#  HELPER FUNCTIONS
# ============================================================================

def set_default_font(doc):
    """Set default font for the entire document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)

    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def set_page_margins(doc):
    """Set A4 page size with specified margins."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)
        section.bottom_margin = Cm(3)


def add_bab_heading(doc, text):
    """Add BAB heading: UPPERCASE, Bold, Centered, 12pt before / 6pt after."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = LINE_SPACING

    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    return p


def add_sub_heading(doc, text):
    """Add sub-heading: Title Case, Bold, Left-aligned, 12pt before / 6pt after."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = LINE_SPACING

    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    return p


def add_subsub_heading(doc, text):
    """Add sub-sub-heading: same as sub_heading."""
    return add_sub_heading(doc, text)


def _apply_justify_wordwrap(p):
    """Apply word-wrap to avoid excessive spacing in justified paragraphs."""
    pPr = p._element.get_or_add_pPr()
    # w:jc already set via p.alignment; add wordWrap to allow mid-word breaks
    wordWrap = pPr.find(qn('w:wordWrap'))
    if wordWrap is None:
        wordWrap = parse_xml(f'<w:wordWrap {nsdecls("w")} w:val="1"/>')
        pPr.append(wordWrap)


def add_body(doc, text, indent=True):
    """Add justified body paragraph with optional first-line indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = FIRST_LINE_INDENT
    _apply_justify_wordwrap(p)

    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    return p


def add_body_mixed(doc, parts, indent=True):
    """Add a paragraph with mixed bold/italic/normal runs.
    parts: list of (text, bold, italic) tuples
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = FIRST_LINE_INDENT
    _apply_justify_wordwrap(p)

    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
            rpr.append(rFonts)
    return p


def add_numbered_item(doc, number, bold_text, rest_text, is_italic_bold=False):
    """Add a numbered paragraph: 'N. **bold_text** rest_text'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)
    _apply_justify_wordwrap(p)

    parts = []
    parts.append((f"{number}. ", False, False))
    parts.append((bold_text, True, is_italic_bold))
    parts.append((f" {rest_text}", False, False))

    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
            rpr.append(rFonts)
    return p


def add_table_caption(doc, text):
    """Add table caption: Bold, centered, above the table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing = SINGLE_SPACING

    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    return p


def add_figure_caption(doc, text):
    """Add figure caption: Bold, centered, below the figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing = SINGLE_SPACING

    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    return p


def add_formula_centered(doc, formula_text, number_text=""):
    """Add a centered formula with optional number on the right."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = LINE_SPACING

    run = p.add_run(formula_text)
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    if number_text:
        run2 = p.add_run(f"    {number_text}")
        run2.font.name = FONT_NAME
        run2.font.size = FONT_SIZE
    return p


def add_figure_placeholder(doc, caption_text):
    """Add placeholder for figure + caption below."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    run = p.add_run("[Gambar disisipkan di sini]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    add_figure_caption(doc, caption_text)


def create_table(doc, headers, rows):
    """Create a formatted table with single spacing."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = SINGLE_SPACING
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        run = p.add_run(header)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        # Shade header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Fill data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.line_spacing = SINGLE_SPACING
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            # Center align for narrow columns, left for wide ones
            if col_idx == 0 or (len(headers) > 2 and col_idx in [0, 2, 3]):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


def add_empty_line(doc):
    """Add an empty line."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = FONT_SIZE
    return p


# ============================================================================
#  MAIN DOCUMENT GENERATION
# ============================================================================

def generate_document():
    doc = Document()
    set_default_font(doc)
    set_page_margins(doc)

    # ========================================================================
    # BAB I PENDAHULUAN
    # ========================================================================
    add_bab_heading(doc, "BAB I PENDAHULUAN")

    add_sub_heading(doc, "1.1 Latar Belakang")

    add_body(doc,
        "Pengelolaan bantuan sosial adalah instrumen krusial pemerintah untuk menekan angka kemiskinan dan "
        "meningkatkan taraf hidup masyarakat prasejahtera. Agar program seperti Program Keluarga Harapan (PKH) "
        "berjalan efektif, penyalurannya harus transparan dan akurat (Altha Inas Shofyana et al., 2025; "
        "Purnomo et al., 2025). Akurasi data menjadi fondasi utama untuk mencapai tujuan kesejahteraan sosial "
        "tanpa adanya kecurangan dari pihak pelaksana (Limbong et al., 2023)."
    )

    add_body(doc,
        "Namun, realita di Kantor Kecamatan Pondok Aren menunjukkan proses seleksi masih terhambat oleh "
        "subjektivitas dan inefisiensi manual. Berdasarkan data kependudukan (Dinas Kependudukan dan Pencatatan "
        "Sipil Kota Tangerang Selatan, 2025), dari sekitar 77.000 KK di 11 kelurahan, terdapat \u00b12.300 KK "
        "prasejahtera. Sementara itu, berdasarkan hasil observasi dan wawancara di instansi, kuota bansos "
        "rata-rata hanya 850 per tahun. Beban validasi yang hanya bertumpu pada 4\u20135 staf menyebabkan "
        "pengambilan keputusan menjadi lambat dan kurang transparan. Penilaian konvensional ini sering luput "
        "memperhatikan kriteria Kementerian Sosial, sehingga memicu risiko distribusi bantuan yang tidak adil "
        "(Altha Inas Shofyana et al., 2025)."
    )

    add_body(doc,
        "Untuk mengatasi masalah tersebut, implementasi Sistem Pendukung Keputusan (SPK) menggunakan metode "
        "Simple Additive Weighting (SAW) menjadi solusi yang mendesak. SPK memfasilitasi komputasi cepat "
        "dengan biaya rendah (Limbong et al., 2023) untuk menghasilkan perankingan yang objektif "
        "(Larantukan et al., 2025). Pemilihan algoritma SAW didasarkan pada keunggulannya dalam mengolah "
        "kriteria kuantitatif dan kualitatif melalui proses normalisasi matriks serta penjumlahan terbobot "
        "(Adhika Pramita Widyassari et al., 2023; Naibaho, 2026). Metode ini secara ilmiah terbukti efektif "
        "memberikan urutan prioritas berdasarkan skor akhir tertinggi (Suprapto et al., 2024)."
    )

    add_body(doc,
        "Sistem SPK tersebut dirancang berbasis web untuk meningkatkan aksesibilitas, akuntabilitas, dan "
        "validitas data seleksi. Platform ini memungkinkan petugas mengolah data secara real-time dan "
        "meminimalisasi risiko kehilangan dokumen fisik. Melalui Kerja Praktik berjudul \"Sistem Pendukung "
        "Keputusan Kelayakan Penerima Bantuan Sosial Menggunakan Metode SAW Berbasis Web pada Kantor "
        "Kecamatan Pondok Aren\", diharapkan terjadi transformasi pelayanan publik dari manual menuju "
        "digital. Dengan demikian, distribusi bantuan sosial dapat dipastikan lebih profesional, transparan, "
        "dan tepat sasaran bagi warga yang paling berhak."
    )

    # 1.2 Identifikasi Masalah
    add_sub_heading(doc, "1.2 Identifikasi Masalah")
    add_body(doc,
        "Berdasarkan uraian latar belakang di atas, permasalahan yang menjadi fokus penelitian ini "
        "dapat diidentifikasi sebagai berikut:"
    )
    add_numbered_item(doc, 1, "Subjektivitas Penilaian.",
        "Verifikasi manual oleh petugas rentan terhadap interpretasi individu, sehingga standar penilaian antarkelurahan menjadi tidak seragam.")
    add_numbered_item(doc, 2, "Inefisiensi Data Manual.",
        "Pengelolaan data konvensional mengakibatkan lambatnya rekapitulasi, tingginya risiko duplikasi, dan sulitnya pembaruan data secara real-time.")
    add_numbered_item(doc, 3, "Ketiadaan Sistem Pendukung Keputusan.",
        "Belum adanya aplikasi perhitungan multikriteria membuat proses seleksi tidak transparan dan sulit dipertanggungjawabkan secara kuantitatif.")

    # 1.3 Tujuan Penulisan
    add_sub_heading(doc, "1.3 Tujuan Penulisan")
    add_body(doc, "Tujuan dari pelaksanaan Kerja Praktik dan penulisan laporan ini adalah:")
    add_numbered_item(doc, 1, "Pengembangan Perangkat Lunak.",
        "Membangun aplikasi SPK berbasis web menggunakan metode SAW untuk menentukan kelayakan penerima bantuan secara otomatis dan objektif.")
    add_numbered_item(doc, 2, "Peningkatan Kualitas Layanan.",
        "Mengoptimalkan efisiensi waktu, akurasi, dan transparansi proses seleksi di Kantor Kecamatan Pondok Aren melalui perankingan otomatis.")
    add_numbered_item(doc, 3, "Penerapan Ilmu Informatika.",
        "Mengimplementasikan solusi metode SAW berbasis web di instansi kecamatan untuk mentransformasi seleksi manual menjadi keputusan digital yang terukur dan dapat dipertanggungjawabkan.")

    # 1.4 Batasan Permasalahan
    add_sub_heading(doc, "1.4 Batasan Permasalahan")
    add_body(doc, "Agar pembahasan tetap fokus dan terarah, batasan permasalahan ditetapkan sebagai berikut:")

    add_subsub_heading(doc, "1.4.1 Batasan Masalah Penelitian")
    add_numbered_item(doc, 1, "Objek dan Lokasi.",
        "Penelitian difokuskan pada pengolahan data warga di 11 kelurahan wilayah Kantor Kecamatan Pondok Aren. Data pengujian merupakan sampel yang disamarkan (anonymized).")
    add_numbered_item(doc, 2, "Kriteria Penilaian.",
        "Parameter didasarkan pada standar indikator kemiskinan instansi dan regulasi Kementerian Sosial.")
    add_numbered_item(doc, 3, "Ruang Lingkup Keputusan.",
        "Sistem berfokus pada perankingan dan penentuan status kelayakan berdasarkan ambang batas nilai preferensi menggunakan algoritma SAW.")

    add_subsub_heading(doc, "1.4.2 Gambaran Sistem Informasi dan Sub-Sistem")
    add_numbered_item(doc, 1, "Platform dan Teknologi.",
        "Aplikasi berbasis web dikembangkan menggunakan Python (framework Flask) dan basis data SQLite agar mudah diakses melalui peramban (browser).")
    add_numbered_item(doc, 2, "Sub-sistem Keputusan.",
        "Mesin komputasi (engine) menggunakan algoritma SAW yang mencakup normalisasi matriks dan penjumlahan terbobot.")
    add_numbered_item(doc, 3, "Fungsionalitas.",
        "Sistem bersifat standalone (tidak tersinkronisasi otomatis dengan pangkalan data kementerian) dengan fitur manajemen data, pembobotan, perhitungan SAW, dan pelaporan.")

    # 1.5 Metode Penelitian
    add_sub_heading(doc, "1.5 Metode Penelitian")
    add_body(doc,
        "Metode penelitian yang digunakan dalam Kerja Praktik ini meliputi dua pendekatan utama, "
        "yaitu pengumpulan data dan pengembangan sistem."
    )

    add_subsub_heading(doc, "1.5.1 Metode Pengumpulan Data")
    add_numbered_item(doc, 1, "Metode Observasi.",
        "Melakukan pengamatan langsung di Kantor Kecamatan Pondok Aren untuk memahami alur birokrasi, prosedur verifikasi, dan kendala teknis pendataan manual.")
    add_numbered_item(doc, 2, "Metode Wawancara.",
        "Melaksanakan diskusi mendalam dengan staf Seksi Kemasyarakatan guna memperoleh informasi mengenai parameter kelayakan, bobot kriteria, dan kebutuhan fungsional sistem. Melalui wawancara ini, disepakati bahwa data penduduk disimulasikan secara valid demi menjaga kerahasiaan dokumen internal instansi.")
    add_numbered_item(doc, 3, "Studi Pustaka.",
        "Mempelajari literatur ilmiah terkait Sistem Pendukung Keputusan, algoritma SAW, dan panduan penulisan laporan dari universitas.")

    add_subsub_heading(doc, "1.5.2 Metode Pengembangan Sistem")
    add_body(doc,
        "Pengembangan sistem menggunakan metode Waterfall (Linear Sequential Model). Metode ini dipilih "
        "karena kebutuhan sistem telah terdefinisi secara jelas sejak awal melalui observasi dan wawancara, "
        "sehingga pengembangan dapat dilakukan secara linear dan terstruktur. Tahapannya meliputi:"
    )
    add_numbered_item(doc, 1, "Analisis Kebutuhan.",
        "Mengidentifikasi permasalahan sistem berjalan, menentukan tujuh kriteria penilaian, dan merumuskan kebutuhan fungsional aplikasi.")
    add_numbered_item(doc, 2, "Perancangan Sistem.",
        "Menyusun arsitektur sistem menggunakan diagram UML, perancangan basis data (ERD), dan antarmuka pengguna (mockup).")
    add_numbered_item(doc, 3, "Implementasi.",
        "Menerjemahkan rancangan ke dalam kode program menggunakan Python (framework Flask), basis data SQLite, dan logika perhitungan SAW.")
    add_numbered_item(doc, 4, "Pengujian.",
        "Memvalidasi fungsionalitas sistem dan keluaran algoritma menggunakan data simulasi dengan ambang batas nilai preferensi kelayakan 0,50.")
    add_numbered_item(doc, 5, "Pemeliharaan.",
        "Merumuskan saran pengembangan lanjutan, seperti integrasi aplikasi SIAK Terpusat dan pengembangan arsitektur multi-user.")

    # 1.6 Sistematika Penulisan
    add_sub_heading(doc, "1.6 Sistematika Penulisan")
    add_body(doc, "Laporan Kerja Praktik ini disusun dengan sistematika sebagai berikut:")
    add_numbered_item(doc, 1, "BAB I PENDAHULUAN.",
        "Menguraikan Latar Belakang Masalah, Identifikasi Masalah, Tujuan Penulisan, Batasan Permasalahan, Metode Penelitian, dan Sistematika Penulisan.")
    add_numbered_item(doc, 2, "BAB II GAMBARAN UMUM INSTANSI.",
        "Berisi profil Kantor Kecamatan Pondok Aren, struktur organisasi, wewenang, dan infrastruktur teknologi informasi.")
    add_numbered_item(doc, 3, "BAB III PEMBAHASAN.",
        "Membahas Tinjauan Pustaka, analisis sistem berjalan, perancangan sistem usulan (UML), basis data, implementasi algoritma SAW, dan antarmuka sistem.")
    add_numbered_item(doc, 4, "BAB IV PENUTUP.",
        "Berisi Kesimpulan atas hasil implementasi sistem dan Saran pengembangan di masa mendatang.")

    # Page break
    doc.add_page_break()

    # ========================================================================
    # BAB II GAMBARAN UMUM INSTANSI
    # ========================================================================
    add_bab_heading(doc, "BAB II GAMBARAN UMUM INSTANSI")

    add_sub_heading(doc, "2.1 Penjelasan Instansi Tempat Kerja Praktik")
    add_body(doc,
        "Kantor Kecamatan Pondok Aren merupakan unsur perangkat daerah yang menjalankan fungsi pelayanan "
        "publik dan pemerintahan umum di tingkat kewilayahan Kota Tangerang Selatan. Instansi ini berkedudukan "
        "di Jalan Graha Bintaro Nomor 1, Kelurahan Perigi Baru, Kecamatan Pondok Aren, Kota Tangerang Selatan."
    )
    add_body(doc,
        "Sebagai salah satu kecamatan dengan kepadatan penduduk tertinggi di Tangerang Selatan, instansi "
        "ini melayani wilayah seluas 29,88 km\u00b2 yang mencakup 11 kelurahan: Perigi Baru, Perigi Lama, "
        "Pondok Kacang Barat, Pondok Kacang Timur, Pondok Pucung, Pondok Jaya, Pondok Aren, Jurang Mangu "
        "Barat, Jurang Mangu Timur, Pondok Karya, dan Pondok Betung. Secara strategis, instansi ini menjadi "
        "simpul penghubung antara kebijakan Pemerintah Kota dengan kebutuhan masyarakat di wilayah yang "
        "berbatasan langsung dengan Jakarta Selatan dan Kota Tangerang."
    )

    # 2.2
    add_sub_heading(doc, "2.2 Sejarah, Struktur Organisasi, Tugas, dan Wewenang")

    add_subsub_heading(doc, "2.2.1 Sejarah Singkat dan Transformasi Wilayah")
    add_body(doc,
        "Nama \"Pondok Aren\" memiliki nilai historis yang kuat, diambil dari kondisi alam masa lampau "
        "di mana wilayah ini merupakan kampung besar yang banyak ditumbuhi pohon aren (Arenga pinnata). "
        "Lintasan sejarah instansi ini dapat dirangkum dalam beberapa tonggak waktu penting sebagai berikut:"
    )
    add_numbered_item(doc, 1, "Tahun 1981\u20131982:",
        "Pembentukan Kecamatan Pondok Aren sebagai hasil pemekaran dari Kecamatan Ciledug saat masih di bawah administrasi Kabupaten Tangerang, Provinsi Jawa Barat.")
    add_numbered_item(doc, 2, "Tahun 1983:",
        "Peresmian gedung pelayanan pertama oleh Bupati Tangerang (H. Tajus Sobirin) di atas lahan eks perkebunan karet PTP XI di wilayah Desa Pondok Aren.")
    add_numbered_item(doc, 3, "Tahun 2005:",
        "Perpindahan kantor ke lokasi saat ini di Kelurahan Perigi Baru karena lokasi kantor lama diklaim oleh pengembang Bintaro Jaya.")
    add_numbered_item(doc, 4, "Tahun 2008\u2013Sekarang:",
        "Integrasi penuh ke dalam wilayah otonom Kota Tangerang Selatan setelah pembentukannya disahkan melalui UU No. 51 Tahun 2008.")

    # 2.2.2 Struktur Organisasi
    add_subsub_heading(doc, "2.2.2 Struktur Organisasi")
    add_body(doc,
        "Struktur organisasi Kecamatan Pondok Aren disusun untuk mendukung pelayanan prima kepada masyarakat. "
        "Pimpinan tertinggi adalah seorang Camat (saat ini dijabat oleh H. Hendra Gunawan, S.H., M.Si.) "
        "yang bertanggung jawab kepada Wali Kota. Secara visual, susunan hierarki dan pembagian tugas dalam "
        "organisasi tersebut dapat diamati pada Gambar 1 berikut."
    )
    add_figure_placeholder(doc, "Gambar 1 \u2013 Struktur Organisasi Kecamatan Pondok Aren")

    add_body(doc,
        "Berdasarkan struktur organisasi di atas, rincian tugas pokok dan wewenang fungsional pada setiap "
        "bagian dijabarkan secara lengkap pada Tabel 1 berikut."
    )

    # Tabel 1
    add_table_caption(doc, "Tabel 1 \u2013 Tugas Pokok dan Wewenang Fungsional Kecamatan Pondok Aren")
    create_table(doc,
        ["Jabatan / Bagian", "Tugas Pokok dan Wewenang"],
        [
            ["Camat", "Mengoordinasikan penyelenggaraan pemerintahan umum, membina kelurahan, serta menjaga ketenteraman dan ketertiban wilayah."],
            ["Sekretaris Camat", "Mengelola administrasi internal yang mencakup tata usaha, kepegawaian, keuangan, dan penyusunan laporan akuntabilitas kinerja."],
            ["Seksi Tata Pemerintahan", "Mengelola administrasi kependudukan, pertanahan, monografi kecamatan, serta pembinaan administrasi kelurahan."],
            ["Seksi Pelayanan Umum", "Menyelenggarakan Pelayanan Administrasi Terpadu Kecamatan (PATEN) serta memproses perizinan dan non-perizinan."],
            ["Seksi Ketenteraman dan Ketertiban", "Melakukan pengawasan penegakan Perda, koordinasi dengan aparat keamanan, dan menjaga ketertiban umum."],
            ["Seksi Ekonomi dan Pembangunan", "Memfasilitasi pembangunan sarana prasarana fisik, pemeliharaan fasilitas umum, dan pemantauan ekonomi lokal."],
            ["Seksi Kemasyarakatan", "Membina lembaga kemasyarakatan (LPM, PKK, Karang Taruna), memfasilitasi pemberdayaan sosial, serta mengelola program bantuan sosial di wilayah kecamatan."],
        ]
    )

    # 2.3
    add_sub_heading(doc, "2.3 Penjelasan Unit Tempat Riset (Seksi Kemasyarakatan)")
    add_body(doc,
        "Riset Kerja Praktik ini difokuskan pada Seksi Kemasyarakatan, yang memiliki tugas pokok membina "
        "lembaga kemasyarakatan dan memfasilitasi pemberdayaan sosial di wilayah Kecamatan Pondok Aren. "
        "Unit ini bertanggung jawab langsung terhadap proses pendataan, verifikasi, dan penentuan prioritas "
        "calon penerima berbagai program bantuan sosial, seperti Bantuan Langsung Tunai (BLT), Program "
        "Keluarga Harapan (PKH), dan Bantuan Pangan Non-Tunai (BPNT)."
    )
    add_body(doc,
        "Kompleksitas penentuan prioritas penerima bantuan yang saat ini masih dilakukan secara manual "
        "mendorong perlunya dukungan Sistem Pendukung Keputusan (SPK) di unit ini. Fokus riset pada Seksi "
        "Kemasyarakatan diambil karena relevansinya yang tinggi dengan tujuan membangun sistem berbasis "
        "algoritma Simple Additive Weighting (SAW), guna menghasilkan keputusan seleksi yang lebih objektif, "
        "transparan, dan terukur."
    )

    # 2.4
    add_sub_heading(doc, "2.4 Infrastruktur Teknologi Informasi")
    add_body(doc,
        "Infrastruktur teknologi informasi (TI) di Kantor Kecamatan Pondok Aren dirancang secara "
        "menyeluruh untuk mendukung implementasi Sistem Pemerintahan Berbasis Elektronik (SPBE)."
    )

    add_subsub_heading(doc, "2.4.1 Jaringan dan Konektivitas")
    add_body(doc,
        "Konektivitas di area kantor memanfaatkan jaringan fiber optic (FO) dengan topologi Local Area "
        "Network (LAN). Instansi ini telah mengimplementasikan protokol IPv6 dan Virtual Private Network "
        "(VPN) yang dikelola oleh Diskominfo Kota Tangerang Selatan. Konfigurasi jaringan ini sangat esensial "
        "karena VPN menjamin keamanan transmisi data internal, sementara IPv6 memastikan stabilitas "
        "skalabilitas aksesibilitas ketika sistem pendukung keputusan (SPK) berbasis web yang dibangun "
        "diakses secara bersamaan oleh staf pelayanan."
    )

    add_subsub_heading(doc, "2.4.2 Perangkat Keras")
    add_body(doc,
        "Perangkat keras (hardware) yang digunakan untuk operasional pelayanan di Kantor Kecamatan Pondok "
        "Aren meliputi empat kategori utama berikut:"
    )
    add_numbered_item(doc, 1, "Komputer Operasional (Client).",
        "Menjalankan sistem operasi Windows 10 dengan spesifikasi prosesor kelas menengah, memori RAM 4\u20138 GB, dan media penyimpanan Hard Disk Drive (HDD) berkapasitas 500 GB hingga 1 TB. Spesifikasi ini sangat memadai untuk mengeksekusi aplikasi SPK berbasis web yang dikembangkan.")
    add_numbered_item(doc, 2, "Server Lokal dan Jaringan.",
        "Menggunakan Router Mikrotik seri RB yang bertindak sebagai firewall utama untuk sistem keamanan data dan manajemen bandwidth, serta switch Gigabit untuk kelancaran distribusi data antarruangan.")
    add_numbered_item(doc, 3, "Perangkat Biometrik.",
        "Terdiri atas pemindai sidik jari, kamera digital (webcam e-KTP), dan iris scanner yang terintegrasi langsung dengan aplikasi SIAK pusat.")
    add_numbered_item(doc, 4, "Perangkat Cetak.",
        "Menggunakan printer laser jet standar untuk pencetakan dokumen di kertas HVS, dan printer thermal khusus untuk mencetak blangko KTP-el.")

    add_subsub_heading(doc, "2.4.3 Perangkat Lunak")
    add_body(doc,
        "Sistem informasi (software) yang digunakan dalam operasional Kantor Kecamatan Pondok Aren "
        "terdiri atas tiga aplikasi utama berikut:"
    )
    add_numbered_item(doc, 1, "SIAK Terpusat.",
        "Aplikasi utama dari Kemendagri untuk pengelolaan database kependudukan secara real-time.")
    add_numbered_item(doc, 2, "SIMPONIE.",
        "Inovasi aplikasi untuk manajemen perizinan online di tingkat kecamatan.")
    add_numbered_item(doc, 3, "Sobat Dukcapil.",
        "Platform pelayanan daring untuk permohonan akta dan kartu identitas secara mandiri.")

    # 2.5
    add_sub_heading(doc, "2.5 Proses Bisnis Instansi")
    add_body(doc,
        "Proses bisnis di Kantor Kecamatan Pondok Aren berfokus pada peningkatan efisiensi layanan "
        "melalui digitalisasi alur kerja."
    )

    add_subsub_heading(doc, "2.5.1 Alur Pelayanan Administrasi (PATEN)")
    add_body(doc,
        "Prosedur standar pelayanan umum mengikuti tahapan yang telah terstruktur, meliputi: penerimaan "
        "berkas, verifikasi kelengkapan melalui sistem SIAK, pemrosesan (entry data atau perekaman biometrik), "
        "pencetakan dokumen, dan penyerahan langsung kepada pemohon."
    )

    add_subsub_heading(doc, "2.5.2 Program Inovatif \"Jemput Bola\"")
    add_body(doc,
        "Selain melayani di kantor, instansi juga menjalankan program SAPA WARGA yang melibatkan tim teknis "
        "untuk mendatangi lingkungan warga secara langsung. Program ini dirancang khusus untuk melakukan "
        "perekaman data kependudukan bagi kelompok rentan yang sulit mengakses layanan kantor, seperti "
        "lansia dan penyandang disabilitas."
    )

    add_subsub_heading(doc, "2.5.3 Alur Verifikasi dan Seleksi Calon Penerima Bantuan Sosial")
    add_body(doc,
        "Proses bisnis yang terkait langsung dengan penelitian ini adalah mekanisme seleksi penerima "
        "bantuan sosial di Seksi Kemasyarakatan. Untuk memberikan gambaran yang komprehensif, tahapan "
        "prosedur konvensional tersebut divisualisasikan dalam bagan alir (flowchart) pada Gambar 2 berikut."
    )
    add_figure_placeholder(doc, "Gambar 2 \u2013 Flowchart Alur Verifikasi dan Seleksi Calon Penerima Bantuan Sosial")

    add_body(doc,
        "Berdasarkan bagan alir di atas, rincian penjelasan dari setiap prosedur administratif yang diterapkan adalah sebagai berikut:"
    )
    add_numbered_item(doc, 1, "Pengajuan Data.",
        "Pihak kelurahan mengirimkan data usulan dalam bentuk dokumen cetak atau spreadsheet. Karena tidak ada format baku, staf kecamatan harus menyesuaikan struktur data secara manual sebelum diproses.")
    add_numbered_item(doc, 2, "Rekapitulasi Manual.",
        "Penggabungan data dari 11 kelurahan dilakukan menggunakan aplikasi perkantoran standar. Proses ini memakan waktu lama dan rentan terhadap kesalahan (human error) seperti duplikasi atau ketidakkonsistenan penulisan nama.")
    add_numbered_item(doc, 3, "Verifikasi Lapangan.",
        "Petugas melakukan kunjungan fisik ke rumah warga untuk memeriksa kondisi tempat tinggal, penghasilan, dan tanggungan. Penilaian ini bersifat subjektif karena murni bergantung pada interpretasi visual petugas tanpa instrumen pengukur yang terstandarisasi.")
    add_numbered_item(doc, 4, "Penetapan Prioritas.",
        "Daftar penerima bantuan ditetapkan melalui musyawarah konsensus internal. Keputusan diambil secara kualitatif, bukan kuantitatif, sehingga sering menyulitkan instansi saat diminta pertanggungjawaban oleh masyarakat mengenai dasar urutan prioritas.")
    add_numbered_item(doc, 5, "Penyaluran.",
        "Daftar penerima tahap akhir dikirimkan kembali ke kelurahan untuk didistribusikan kepada warga yang berhak.")

    add_body(doc,
        "Kelemahan utama dari alur kerja ini adalah ketiadaan mekanisme perhitungan pembobotan yang objektif. "
        "Oleh karena itu, penerapan Sistem Pendukung Keputusan menggunakan algoritma SAW dalam Kerja Praktik "
        "ini dirancang untuk menggantikan musyawarah subjektif tersebut dengan kalkulasi sistem yang akurat "
        "dan dapat dipertanggungjawabkan."
    )

    # 2.6 Tinjauan Pustaka
    add_sub_heading(doc, "2.6 Tinjauan Pustaka")

    add_subsub_heading(doc, "2.6.1 Sistem Pendukung Keputusan (SPK)")
    add_subsub_heading(doc, "a. Definisi Menurut Para Ahli")
    add_body(doc,
        "Sistem Pendukung Keputusan (SPK) didefinisikan sebagai kerangka kerja terkomputerisasi yang "
        "dirancang untuk membantu individu atau organisasi dalam memecahkan masalah semi-terstruktur "
        "dengan memanfaatkan pengolahan data dan model matematis. Sejalan dengan hal tersebut, SPK "
        "dirancang untuk membantu otoritas dalam mempertimbangkan banyak kriteria yang kompleks secara "
        "objektif (Altha Inas Shofyana et al., 2025). Sistem ini berfungsi sebagai sistem informasi "
        "interaktif yang menyediakan pemodelan dan manipulasi data untuk mendukung pengambilan keputusan "
        "yang lebih berkualitas (Limbong et al., 2023)."
    )

    add_subsub_heading(doc, "b. Karakteristik dan Komponen SPK")
    add_body(doc,
        "Karakteristik utama SPK adalah kemampuannya dalam melakukan komputasi data dalam jumlah besar "
        "secara cepat dengan biaya operasional yang relatif rendah (Limbong et al., 2023). Komponen SPK "
        "umumnya terdiri atas modul manajemen kriteria untuk penyesuaian bobot, modul masukan alternatif "
        "data warga, serta mesin perhitungan algoritma yang menghasilkan keluaran objektif berupa nilai "
        "preferensi (Altha Inas Shofyana et al., 2025)."
    )

    add_subsub_heading(doc, "c. Tujuan SPK dalam Pengambilan Keputusan")
    add_body(doc,
        "Tujuan utama SPK bukan untuk menggantikan peran otoritas pengambil keputusan, melainkan untuk "
        "meningkatkan efektivitas keputusan yang diambil (Limbong et al., 2023). Dalam konteks tata kelola "
        "di Kecamatan Pondok Aren, SPK bertujuan mengubah data mentah kondisi sosial-ekonomi masyarakat "
        "menjadi landasan komputasi ilmiah yang akuntabel, sehingga membantu pihak kecamatan menentukan "
        "prioritas penerima bantuan secara tepat sasaran."
    )

    add_subsub_heading(doc, "d. Proses Pengambilan Keputusan")
    add_body(doc,
        "Proses pengambilan keputusan dalam SPK melibatkan identifikasi masalah, penentuan kriteria "
        "penilaian, pencarian alternatif solusi, hingga evaluasi menggunakan model matematis "
        "(Altha Inas Shofyana et al., 2025). Proses ini memungkinkan konversi data kualitatif yang bersifat "
        "subjektif dari hasil observasi lapangan menjadi nilai kuantitatif yang terukur dan objektif "
        "(Limbong et al., 2023)."
    )

    # 2.6.2
    add_subsub_heading(doc, "2.6.2 Metode Simple Additive Weighting (SAW)")

    add_subsub_heading(doc, "a. Konsep Dasar SAW")
    add_body(doc,
        "Metode SAW, yang lazim dikenal sebagai metode penjumlahan terbobot, merupakan salah satu metode "
        "pengambilan keputusan multikriteria (Multi-Criteria Decision Making/MCDM) yang banyak digunakan "
        "karena kesederhanaan dan ketepatannya (Altha Inas Shofyana et al., 2025). Konsep dasar metode ini "
        "adalah mencari penjumlahan terbobot dari rating kinerja pada setiap alternatif di seluruh atribut. "
        "Keunggulannya terletak pada efisiensi waktu pemrosesan komputasi serta kemampuannya dalam "
        "menghasilkan urutan prioritas yang transparan (Adhika Pramita Widyassari et al., 2023)."
    )

    add_subsub_heading(doc, "b. Langkah-Langkah Perhitungan")
    add_body(doc,
        "Berdasarkan literatur (Altha Inas Shofyana et al., 2025; Naibaho, 2026), kalkulasi SAW dilakukan "
        "melalui dua tahap utama, yaitu normalisasi matriks dan penjumlahan terbobot."
    )

    add_body_mixed(doc, [("Tahap 1: Normalisasi Matriks", True, False)], indent=False)
    add_body(doc,
        "Normalisasi bertujuan untuk menyamakan skala penilaian setiap kriteria karena masing-masing "
        "kriteria memiliki satuan dan rentang nilai yang berbeda. Rumus normalisasi dibedakan berdasarkan "
        "tipe kriteria sebagai berikut:"
    )
    add_body(doc, "Kriteria Benefit (Keuntungan): Digunakan apabila nilai yang lebih besar menunjukkan kondisi yang lebih baik.", indent=False)
    add_formula_centered(doc, "r\u1d62\u2c7c = x\u1d62\u2c7c / max(x\u1d62\u2c7c)", "(2.1)")
    add_body(doc, "Kriteria Cost (Biaya): Digunakan apabila nilai yang lebih kecil menunjukkan kondisi yang lebih baik.", indent=False)
    add_formula_centered(doc, "r\u1d62\u2c7c = min(x\u1d62\u2c7c) / x\u1d62\u2c7c", "(2.2)")

    add_body_mixed(doc, [("Tahap 2: Penjumlahan Terbobot", True, False)], indent=False)
    add_body(doc,
        "Setelah normalisasi selesai dilakukan, nilai preferensi akhir (V\u1d62) untuk setiap alternatif "
        "dihitung dengan menjumlahkan hasil perkalian antara bobot kriteria (w\u2c7c) dan nilai ternormalisasi "
        "(r\u1d62\u2c7c), sebagaimana dirumuskan berikut:"
    )
    add_formula_centered(doc, "V\u1d62 = \u03a3 w\u2c7c \u00d7 r\u1d62\u2c7c", "(2.3)")

    # 2.6.3
    add_subsub_heading(doc, "2.6.3 Aplikasi Berbasis Web")

    # SYNCHRONIZED: Updated to match actual tech stack from code
    add_subsub_heading(doc, "a. Pengertian Aplikasi Web dan Arsitektur Client-Server")
    add_body(doc,
        "Aplikasi berbasis web beroperasi menggunakan model arsitektur client-server, di mana browser "
        "klien mengirimkan permintaan (request) melalui jaringan dan server memberikan respons berupa "
        "data yang kemudian ditampilkan kepada pengguna."
    )

    add_subsub_heading(doc, "b. Pengembangan Sisi Depan (Frontend)")
    add_body(doc,
        "Antarmuka pengguna (frontend) sistem dikembangkan menggunakan HTML sebagai kerangka struktur "
        "halaman, Tailwind CSS (CLI/Local Build) untuk penataan tampilan, dan JavaScript untuk "
        "interaktivitas antarmuka. Selain itu, pustaka Chart.js digunakan untuk merender grafik statistik "
        "pada halaman Dashboard."
    )

    add_subsub_heading(doc, "c. Pengembangan Sisi Belakang (Backend)")
    add_body(doc,
        "Sisi belakang (backend) sistem menangani seluruh logika bisnis menggunakan Python dengan "
        "framework Flask, templating engine Jinja2 untuk penghasil halaman dinamis, dan Object-Relational "
        "Mapping (ORM) SQLAlchemy untuk pengelolaan basis data."
    )

    add_subsub_heading(doc, "d. Basis Data SQLite")
    add_body(doc,
        "SQLite dipilih sebagai sistem manajemen basis data karena sifatnya yang serverless dan tidak "
        "memerlukan proses instalasi terpisah. Karakteristik ini sangat ideal untuk diimplementasikan di "
        "infrastruktur instansi kecamatan yang memerlukan sistem operasional secara standalone yang praktis "
        "dan mudah dipelihara."
    )

    # 2.6.4
    add_subsub_heading(doc, "2.6.4 Penelitian Terdahulu")
    add_body(doc,
        "Untuk memperkuat landasan ilmiah penelitian ini, dilakukan kajian terhadap sejumlah penelitian "
        "terdahulu yang relevan. Hasil pemetaan penelitian tersebut \u2014 mencakup karya Widyassari (2023), "
        "Limbong (2023), Sudi (2024), Suprapto (2024), Shofyana (2025), Larantukan (2025), Muarif (2025), "
        "Purnomo (2025), dan Naibaho (2026) \u2014 dirangkum dalam Tabel 2 berikut."
    )
    add_body(doc, "(Tabel 2 \u2013 Penelitian Terdahulu)", indent=False)

    add_body_mixed(doc, [("Analisis Perbandingan Penelitian Terdahulu", True, False)], indent=False)
    add_body(doc,
        "Berdasarkan kajian terhadap penelitian-penelitian di atas, kebaruan (novelty) dari penelitian "
        "Kerja Praktik ini terletak pada tiga aspek utama:"
    )
    add_numbered_item(doc, 1, "Pendekatan Arsitektur Teknologi.",
        "Menggunakan microframework Flask (Python) dan basis data serverless SQLite, yang membedakannya dari mayoritas penelitian terdahulu yang umumnya menggunakan PHP atau framework berbasis Java.")
    add_numbered_item(doc, 2, "Kustomisasi Kriteria Berbasis Empiris.",
        "Tujuh kriteria sosial-ekonomi yang digunakan dikalibrasi secara spesifik berdasarkan hasil wawancara dan kondisi demografis nyata Kecamatan Pondok Aren, bukan hanya mengadopsi kriteria generik.")
    add_numbered_item(doc, 3, "Otomatisasi Pemrosesan Kelompok (Batch Processing).",
        "Sistem dilengkapi fitur klasifikasi data massal menggunakan pustaka Pandas, yang memungkinkan petugas memproses ratusan data warga sekaligus melalui satu berkas spreadsheet.")

    # Page break
    doc.add_page_break()

    # ========================================================================
    # BAB III PEMBAHASAN
    # ========================================================================
    add_bab_heading(doc, "BAB III PEMBAHASAN")

    add_sub_heading(doc, "3.1 Prosedur Kerja Praktik")
    add_subsub_heading(doc, "3.1.1 Perancangan Sistem")

    add_subsub_heading(doc, "a. Analisis Sistem Berjalan")
    add_body(doc,
        "Prosedur saat ini berjalan melalui pengajuan data dari 11 kelurahan, rekapitulasi manual, "
        "verifikasi lapangan, dan musyawarah penetapan prioritas. Kelemahan fatalnya adalah tingginya "
        "tingkat subjektivitas dalam setiap tahapan penilaian dan ketiadaan hasil akhir berupa perankingan "
        "matematis yang dapat dipertanggungjawabkan secara kuantitatif."
    )

    add_subsub_heading(doc, "b. Analisis Sistem Usulan")
    add_body(doc,
        "Sistem usulan dirancang untuk mentransformasi proses konvensional menjadi digital melalui "
        "digitalisasi data terpusat dan penetapan pembobotan kriteria yang baku. Sebagai dasar perhitungan, "
        "tujuh kriteria penilaian beserta tipe dan bobot masing-masing, yang disepakati melalui wawancara "
        "dengan instansi, dirangkum dalam Tabel 3 berikut."
    )

    # Tabel 3 — SYNCHRONIZED with actual seed data in app.py seed_kriteria()
    add_table_caption(doc, "Tabel 3 \u2013 Bobot Kriteria Sistem Usulan")
    create_table(doc,
        ["Kode", "Nama Kriteria", "Tipe", "Bobot"],
        [
            ["C1", "Penghasilan", "Cost", "0,25"],
            ["C2", "Jumlah Tanggungan", "Benefit", "0,20"],
            ["C3", "Kepemilikan Aset", "Cost", "0,15"],
            ["C4", "Status Rumah", "Cost", "0,10"],
            ["C5", "Kondisi Bangunan", "Cost", "0,10"],
            ["C6", "Daya Listrik", "Cost", "0,10"],
            ["C7", "Sumber Air", "Cost", "0,10"],
            ["Total", "", "", "1,00"],
        ]
    )

    add_subsub_heading(doc, "c. Keunggulan Sistem Usulan")
    add_body(doc,
        "Dibandingkan dengan proses manual yang berjalan, sistem usulan menawarkan sejumlah keunggulan "
        "signifikan, yaitu: objektivitas penilaian berbasis perhitungan matematis, kecepatan pemrosesan "
        "data, transparansi hasil perankingan, akuntabilitas yang dapat diaudit, aksesibilitas melalui "
        "browser, serta potensi integrasi dengan sistem data kependudukan."
    )

    add_subsub_heading(doc, "d. Activity Diagram Sistem Berjalan")
    add_body(doc,
        "Untuk memberikan gambaran yang lebih jelas mengenai alur kerja sistem manual yang sedang berjalan, "
        "berikut disajikan activity diagram pada Gambar 3."
    )
    add_figure_placeholder(doc, "Gambar 3 \u2013 Activity Diagram Sistem Berjalan")

    add_subsub_heading(doc, "e. Activity Diagram Sistem Usulan")
    add_body(doc,
        "Sebagai perbandingan, activity diagram sistem usulan menggambarkan alur kerja yang lebih "
        "terstruktur dan efisien. Terdapat delapan activity diagram yang masing-masing mewakili "
        "satu skenario utama penggunaan sistem, yaitu sebagai berikut."
    )

    # AD 1 — Login
    add_body(doc,
        "Activity diagram pertama menggambarkan proses login Admin ke dalam sistem. Alur dimulai "
        "dari Admin membuka halaman login, memasukkan username dan password, kemudian sistem "
        "memverifikasi kredensial ke database. Jika valid, sesi disimpan dan Admin diarahkan ke "
        "halaman Dashboard; jika tidak valid, pesan error ditampilkan dan Admin diminta mengulang "
        "input."
    )
    add_figure_placeholder(doc, "Gambar 4 \u2013 Activity Diagram Proses Login")

    # AD 2 — Klasifikasi Manual
    add_body(doc,
        "Activity diagram kedua mengilustrasikan proses klasifikasi data warga secara manual. Admin "
        "mengisi form identitas warga (NIK, No KK, nama, pekerjaan, alamat, kelurahan) beserta "
        "nilai tujuh kriteria (C1\u2013C7). Setelah tombol \"Klasifikasikan\" diklik, sistem "
        "memvalidasi format NIK/KK, mengambil data baseline SAW dari database, menghitung skor "
        "SAW, menetapkan status kelayakan berdasarkan threshold \u22650,50, dan menyimpan hasilnya."
    )
    add_figure_placeholder(doc, "Gambar 5 \u2013 Activity Diagram Proses Klasifikasi Manual")

    # AD 3 — Import Massal
    add_body(doc,
        "Activity diagram ketiga menggambarkan proses import massal (batch) data warga melalui "
        "berkas Excel atau CSV. Admin memilih tab \"Import Massal\", memilih file, lalu mengunggahnya. "
        "Sistem memvalidasi format file (.csv/.xls/.xlsx) dan isi setiap baris data. Jika seluruh "
        "baris valid, sistem memetakan nilai kolom ke sub-kriteria, menghitung SAW untuk setiap "
        "baris secara iteratif, dan menyimpan semua record ke database dalam satu batch commit."
    )
    add_figure_placeholder(doc, "Gambar 6 \u2013 Activity Diagram Proses Import Massal (Batch)")

    # AD 4 — Edit Histori
    add_body(doc,
        "Activity diagram keempat mengilustrasikan proses pengeditan data hasil klasifikasi yang "
        "telah tersimpan. Admin membuka halaman Histori, memilih record yang akan diedit, lalu "
        "mengubah data identitas dan/atau nilai kriteria. Secara opsional, Admin dapat memilih "
        "override status kelayakan secara manual. Sistem menghitung ulang skor SAW berdasarkan data "
        "yang diperbarui dan menyimpan perubahan ke database."
    )
    add_figure_placeholder(doc, "Gambar 7 \u2013 Activity Diagram Proses Edit Data Histori")

    # AD 5 — Hapus Data
    add_body(doc,
        "Activity diagram kelima menggambarkan proses penghapusan data histori. Admin dapat "
        "menghapus satu record terpilih atau seluruh data histori sekaligus. Untuk penghapusan "
        "massal, sistem menampilkan dialog konfirmasi terlebih dahulu guna mencegah penghapusan "
        "yang tidak disengaja. Setelah dikonfirmasi, sistem menghapus record dari database dan "
        "menampilkan pesan sukses."
    )
    add_figure_placeholder(doc, "Gambar 8 \u2013 Activity Diagram Proses Hapus Data Histori")

    # AD 6 — Filter & Ekspor
    add_body(doc,
        "Activity diagram keenam mengilustrasikan proses filter dan ekspor data histori. Admin "
        "membuka halaman Histori, memilih filter berdasarkan kelurahan dan/atau status kelayakan, "
        "lalu mengklik tombol \"Ekspor ke Excel\". Sistem mengambil data yang telah difilter, "
        "mengonversinya ke format Excel menggunakan pustaka openpyxl melalui Pandas, dan "
        "mengirimkan file sebagai unduhan ke browser Admin."
    )
    add_figure_placeholder(doc, "Gambar 9 \u2013 Activity Diagram Proses Filter dan Ekspor Data")

    # AD 7 — Manajemen Kriteria
    add_body(doc,
        "Activity diagram ketujuh menggambarkan proses manajemen kriteria dan sub-kriteria. Admin "
        "membuka menu Kriteria dan dapat memilih aksi tambah, edit, atau hapus kriteria. Untuk "
        "aksi tambah dan edit, sistem memvalidasi total bobot agar tidak melebihi 1,00 sebelum "
        "menyimpan perubahan ke database. Jika validasi gagal, pesan error ditampilkan dan Admin "
        "diminta memperbaiki nilai bobot."
    )
    add_figure_placeholder(doc, "Gambar 10 \u2013 Activity Diagram Proses Manajemen Kriteria")

    # AD 8 — Manajemen Profil
    add_body(doc,
        "Activity diagram kedelapan mengilustrasikan proses manajemen profil dan penggantian "
        "password Admin. Admin dapat memilih antara mengedit profil (mengubah nama lengkap "
        "dan/atau foto profil) atau mengganti password (mengisi password lama, password baru, dan "
        "konfirmasi). Sistem memvalidasi seluruh input sebelum menyimpan perubahan ke database dan "
        "memperbarui data sesi yang aktif."
    )
    add_figure_placeholder(doc, "Gambar 11 \u2013 Activity Diagram Proses Manajemen Profil Admin")

    add_subsub_heading(doc, "f. Normalisasi")
    add_body(doc,
        "Pada implementasi sistem ini, baik kriteria bertipe benefit maupun cost menggunakan bentuk "
        "normalisasi yang sama, sebagaimana ditunjukkan pada persamaan berikut:"
    )
    add_formula_centered(doc, "r\u1d62\u2c7c = x\u1d62\u2c7c / max(x\u1d62\u2c7c)", "(3.1)")
    add_body(doc,
        "Penyeragaman rumus ini dimungkinkan karena skor untuk kriteria cost telah dibalik sejak tahap "
        "pendefinisian sub-kriteria, sehingga kondisi yang semakin membutuhkan bantuan akan memperoleh "
        "skor yang semakin tinggi dengan nilai maksimal 5."
    )

    # SYNCHRONIZED: ERD description matches actual database models
    add_subsub_heading(doc, "g. Entity Relationship Diagram (ERD)")
    add_body(doc,
        "Hubungan antar entitas dalam basis data sistem diilustrasikan melalui Entity Relationship "
        "Diagram (ERD) pada Gambar 12 berikut. Diagram ini menggambarkan bagaimana data pengguna, "
        "kriteria, dan hasil klasifikasi saling berelasi."
    )
    add_figure_placeholder(doc, "Gambar 12 \u2013 Entity Relationship Diagram (ERD)")

    # SYNCHRONIZED: Updated entity descriptions to match actual ORM models
    add_body(doc,
        "Relasi antar entitas yang terbentuk adalah sebagai berikut: tabel users berelasi "
        "satu-ke-banyak (1:N) dengan login_history, serta tabel kriteria berelasi satu-ke-banyak "
        "(1:N) dengan sub_kriteria. Adapun tabel classification_results menyimpan rekam jejak "
        "detail skor kriteria dalam format JSON melalui kolom kriteria_details."
    )

    add_subsub_heading(doc, "h. Sequence Diagram")
    add_body(doc,
        "Interaksi antara pengguna (Admin), antarmuka browser, backend Flask, dan basis data "
        "SQLite untuk setiap skenario utama sistem digambarkan melalui sequence diagram. Terdapat "
        "enam sequence diagram yang masing-masing mewakili satu alur interaksi kritis sebagai "
        "berikut."
    )

    # SD 1 — Login
    add_body(doc,
        "Sequence diagram pertama menggambarkan alur interaksi proses login. Admin mengirimkan "
        "permintaan POST /login dari browser ke Flask backend. Backend memverifikasi kredensial "
        "ke SQLite; jika valid, sesi dibuat, riwayat login dicatat, lalu browser diarahkan ke "
        "halaman Dashboard melalui redirect HTTP 302."
    )
    add_figure_placeholder(doc, "Gambar 13 \u2013 Sequence Diagram Proses Login Admin")

    # SD 2 — Klasifikasi Manual
    add_body(doc,
        "Sequence diagram kedua mengilustrasikan interaksi proses klasifikasi manual. Setelah "
        "Admin mengirim data form (POST /classify), Flask memanggil fungsi hitung_saw() di modul "
        "spk.py, yang mengambil data baseline dari SQLite, menghitung normalisasi matriks dan "
        "penjumlahan terbobot, lalu mengembalikan skor dan status kelayakan. Hasilnya disimpan ke "
        "tabel classification_results dan browser dialihkan ke halaman Histori."
    )
    add_figure_placeholder(doc, "Gambar 14 \u2013 Sequence Diagram Proses Klasifikasi Manual")

    # SD 3 — Import Batch
    add_body(doc,
        "Sequence diagram ketiga menggambarkan interaksi proses import massal. Admin mengunggah "
        "file melalui browser (POST /import). Flask memanggil pustaka Pandas untuk membaca file "
        "secara in-memory menggunakan read_excel() atau read_csv(). Dalam iterasi setiap baris, "
        "sistem memanggil hitung_saw() dan menyimpan hasil ke SQLite. Pada akhirnya, Flask "
        "mengembalikan JSON sukses atau pesan error ke browser."
    )
    add_figure_placeholder(doc, "Gambar 15 \u2013 Sequence Diagram Proses Import Massal (Batch)")

    # SD 4 — Filter & Ekspor
    add_body(doc,
        "Sequence diagram keempat mengilustrasikan interaksi proses filter dan ekspor data. Admin "
        "mengirimkan permintaan GET /export dengan parameter filter ke Flask. Backend mengeksekusi "
        "query berfilter ke SQLite, mengonversi hasil ke DataFrame Pandas, lalu mengekspor ke "
        "format Excel menggunakan openpyxl. File Excel dikirimkan sebagai respons dengan header "
        "Content-Disposition: attachment ke browser Admin."
    )
    add_figure_placeholder(doc, "Gambar 16 \u2013 Sequence Diagram Proses Filter dan Ekspor Data")

    # SD 5 — Edit Histori
    add_body(doc,
        "Sequence diagram kelima menggambarkan interaksi proses edit data histori. Admin mengirim "
        "data form yang telah diperbarui (POST /edit/<id>). Flask mengambil record dari SQLite, "
        "menjalankan ulang kalkulasi hitung_saw() dengan data baru, memeriksa adanya override "
        "status manual, kemudian memperbarui record di database dan mengarahkan kembali ke "
        "halaman Histori."
    )
    add_figure_placeholder(doc, "Gambar 17 \u2013 Sequence Diagram Proses Edit Data Histori")

    # SD 6 — Manajemen Kriteria
    add_body(doc,
        "Sequence diagram keenam mengilustrasikan interaksi proses manajemen kriteria. Admin "
        "mengirimkan form tambah atau edit kriteria (POST /kriteria). Flask memvalidasi total bobot "
        "seluruh kriteria aktif; jika total melebihi 1,00, respons error dikembalikan ke browser. "
        "Jika valid, perubahan disimpan ke tabel kriteria dan sub_kriteria di SQLite, lalu "
        "halaman Manajemen Kriteria diperbarui dengan data terkini."
    )
    add_figure_placeholder(doc, "Gambar 18 \u2013 Sequence Diagram Proses Manajemen Kriteria")

    # 3.1.2
    add_subsub_heading(doc, "3.1.2 Perancangan Perangkat Lunak: Flowchart")
    add_body(doc,
        "Alur kerja perangkat lunak secara keseluruhan \u2014 mulai dari penerimaan input data hingga "
        "penyimpanan hasil ke basis data \u2014 digambarkan dalam flowchart pada Gambar 19 berikut."
    )
    add_figure_placeholder(doc, "Gambar 19 \u2013 Flowchart Proses Klasifikasi Data Warga")
    add_body(doc,
        "Berdasarkan flowchart tersebut, perangkat lunak dirancang untuk mendukung dua mode input "
        "(manual dan massal), mencakup proses validasi data, perhitungan SAW, penentuan status "
        "kelayakan, penyimpanan ke basis data, serta penayangan hasil pada halaman histori."
    )

    # 3.2
    add_sub_heading(doc, "3.2 Analisis dan Pembahasan")

    # SYNCHRONIZED: Description matches actual spk.py module
    add_subsub_heading(doc, "3.2.1 Pembahasan Algoritma")
    add_body(doc,
        "Implementasi algoritma SAW dikemas di dalam modul spk.py sebagai komponen inti sistem. Modul ini "
        "berisi tiga fungsi utama: hitung_saw() untuk kalkulasi normalisasi dan penjumlahan terbobot, "
        "tentukan_kelayakan() untuk menetapkan status berdasarkan ambang batas (threshold) yang "
        "dikonfigurasi di config.py melalui variabel THRESHOLD_LAYAK, dan generate_alasan_dinamis() "
        "untuk menghasilkan penjelasan otomatis atas hasil klasifikasi. Nilai kriteria yang bermakna "
        "\"biaya terkecil\" \u2014 contohnya penghasilan terendah pada kriteria C1 \u2014 langsung dikonversi "
        "menjadi skor prioritas terbesar (5) sejak tahap pendefinisian sub-kriteria. Dengan pendekatan ini, "
        "keseluruhan proses normalisasi cukup menggunakan satu rumus tunggal (nilai / max_skor) yang "
        "berlaku untuk semua kriteria, sehingga logika komputasi menjadi lebih efisien dan mudah diaudit."
    )

    # SYNCHRONIZED: Rancangan Layar matches actual templates
    add_subsub_heading(doc, "3.2.2 Rancangan Layar")
    add_body(doc,
        "Antarmuka sistem dirancang dengan mengutamakan kemudahan penggunaan oleh petugas kecamatan. "
        "Berikut adalah deskripsi singkat setiap halaman yang diimplementasikan dalam sistem:"
    )
    add_numbered_item(doc, 1, "Halaman Login (index.html)",
        "\u2014 Halaman autentikasi untuk memverifikasi identitas Admin sebelum mengakses sistem.")
    add_numbered_item(doc, 2, "Halaman Dashboard (dashboard.html)",
        "\u2014 Halaman utama yang menampilkan ringkasan statistik data warga, grafik distribusi kelayakan (Pie Chart), dan perbandingan per kelurahan (Bar Chart) menggunakan pustaka Chart.js.")
    add_numbered_item(doc, 3, "Halaman Klasifikasi Data (classification.html)",
        "\u2014 Halaman untuk memasukkan data warga secara manual maupun melalui import berkas massal (CSV/XLS/XLSX).")
    add_numbered_item(doc, 4, "Halaman Histori (history.html)",
        "\u2014 Halaman yang menampilkan seluruh rekam hasil klasifikasi yang pernah dilakukan beserta detailnya, dengan fitur pencarian dan filter per kelurahan.")
    add_numbered_item(doc, 5, "Halaman Manajemen Kriteria (kriteria.html)",
        "\u2014 Halaman untuk mengelola kriteria dan bobot penilaian secara dinamis, termasuk manajemen sub-kriteria (sub_kriteria.html).")
    add_numbered_item(doc, 6, "Halaman Informasi SPK (informasi.html)",
        "\u2014 Halaman yang memuat penjelasan tentang metode SAW dan cara kerja sistem.")
    add_numbered_item(doc, 7, "Halaman Profil (profil.html) dan Riwayat Login (riwayat_login.html)",
        "\u2014 Halaman pengelolaan data akun (termasuk foto profil dan nama lengkap) dan rekam jejak akses sistem oleh pengguna. Halaman Ganti Password (ganti_password.html) juga tersedia untuk keamanan akun.")
    add_numbered_item(doc, 8, "Halaman Edit Histori (edit_history.html)",
        "\u2014 Halaman untuk mengedit data hasil klasifikasi yang tersimpan, menghitung ulang skor SAW, dan melakukan override status kelayakan secara manual.")

    # 3.2.3
    add_subsub_heading(doc, "3.2.3 Implementasi dan Integrasi Alur Data")
    add_body(doc,
        "Sistem mengintegrasikan routing Flask dengan antarmuka berbasis Jinja2 untuk menghasilkan respons "
        "halaman yang dinamis. Seluruh logika routing dan CRUD data terdapat di dalam file app.py yang "
        "menjadi entry point aplikasi. Pemrosesan dokumen spreadsheet yang diunggah pengguna dieksekusi "
        "secara instan di memori (in-memory processing) melalui pustaka Pandas, tanpa perlu menyimpan "
        "berkas sementara ke server, sehingga proses lebih cepat dan aman."
    )

    # 3.2.4
    add_subsub_heading(doc, "3.2.4 Pengujian Sistem")
    add_body(doc,
        "Pengujian sistem dilakukan dalam dua tahap untuk memastikan kebenaran algoritma sekaligus "
        "keandalan fungsionalitas aplikasi secara menyeluruh."
    )

    add_subsub_heading(doc, "a. Pengujian Algoritma SAW (Uji Coba dengan Contoh Data)")
    add_body(doc,
        "Untuk memvalidasi keakuratan algoritma, dilakukan uji coba terhadap lima data sampel warga. "
        "Hasil perhitungan SAW beserta status kelayakan masing-masing warga berdasarkan threshold 0,50 "
        "dapat dilihat pada Tabel 5 berikut."
    )

    add_table_caption(doc, "Tabel 5 \u2013 Hasil Uji Coba Program")
    create_table(doc,
        ["No", "Nama", "Kelurahan", "Skor SAW", "Status"],
        [
            ["1", "Dewi Aminah", "Pondok Jaya", "0,92", "Layak"],
            ["2", "Hadi Lestari", "Perigi Lama", "0,90", "Layak"],
            ["3", "Lestari Setiawan", "Pondok Aren", "0,50", "Layak"],
            ["4", "Ayu Lestari", "Perigi Baru", "0,27", "Tidak Layak"],
            ["5", "Tri Setiawan", "Pondok Aren", "0,24", "Tidak Layak"],
        ]
    )

    add_body(doc, "Perhitungan manual algoritma SAW terhadap data sampel di atas dilakukan sebagai berikut:")
    add_body(doc, "Vektor Bobot (W): W = [0,25; 0,20; 0,15; 0,10; 0,10; 0,10; 0,10]", indent=False)
    add_body(doc,
        "Matriks Keputusan (X) dan Normalisasi (R): Karena inversi logika data pada kriteria cost, "
        "normalisasi dilakukan secara linear terhadap nilai maksimum (5) untuk semua kriteria.", indent=False)
    add_body(doc, "Penjumlahan Terbobot (V\u1d62): Nilai preferensi akhir masing-masing alternatif adalah:", indent=False)

    add_formula_centered(doc, "V\u2081 (Dewi Aminah) = 0,92")
    add_formula_centered(doc, "V\u2082 (Hadi Lestari) = 0,90")
    add_formula_centered(doc, "V\u2083 (Lestari Setiawan) = 0,50")
    add_formula_centered(doc, "V\u2084 (Ayu Lestari) = 0,27")
    add_formula_centered(doc, "V\u2085 (Tri Setiawan) = 0,24")

    add_body(doc,
        "Berdasarkan hasil di atas, dengan penetapan threshold sebesar 0,50, sistem berhasil menentukan "
        "status kelayakan seluruh data sampel dengan akurasi logika matematis 100%, sesuai dengan "
        "ekspektasi yang diharapkan."
    )

    add_subsub_heading(doc, "b. Pengujian Fungsional (Black Box Testing)")
    add_body(doc,
        "Pengujian fungsional dilakukan menggunakan teknik Equivalence Partitioning pada empat modul "
        "utama, yaitu Autentikasi, Klasifikasi (Manual dan Massal), Manajemen Kriteria, serta Histori "
        "dan Ekspor. Seluruh skenario pengujian menghasilkan status Valid dan tidak ditemukan bug pada "
        "logika bisnis sistem, sehingga aplikasi dinyatakan lulus pengujian fungsional."
    )

    # 3.2.5
    add_subsub_heading(doc, "3.2.5 Penggunaan Program (User Guide)")
    add_body(doc,
        "Sistem ini dioperasikan oleh Admin Kecamatan melalui serangkaian langkah yang intuitif. Alur "
        "penggunaan dimulai dari proses Login untuk masuk ke sistem, dilanjutkan dengan pembukaan Dashboard "
        "untuk memantau ringkasan data. Selanjutnya, petugas dapat melakukan Klasifikasi Data warga secara "
        "Manual maupun melalui fitur Import massal. Hasil klasifikasi dapat ditinjau pada halaman Histori "
        "dan diekspor ke format Excel untuk keperluan pelaporan. Pengelolaan bobot dan kriteria penilaian "
        "dilakukan melalui halaman Manajemen Kriteria, sementara pengaturan akun dan keamanan dikelola "
        "melalui halaman Profil."
    )

    # Page break
    doc.add_page_break()

    # ========================================================================
    # BAB IV PENUTUP
    # ========================================================================
    add_bab_heading(doc, "BAB IV PENUTUP")

    add_sub_heading(doc, "4.1 Kesimpulan")
    # SYNCHRONIZED: Conclusion matches actual tech stack
    add_body(doc,
        "Berdasarkan seluruh tahapan Kerja Praktik yang telah dilaksanakan di Kantor Kecamatan Pondok "
        "Aren, dapat disimpulkan bahwa Sistem Pendukung Keputusan (SPK) kelayakan penerima bantuan sosial "
        "berbasis web telah berhasil dirancang dan dibangun menggunakan Python (Flask, Jinja2, SQLAlchemy) "
        "dengan basis data SQLite. Implementasi algoritma Simple Additive Weighting (SAW) terbukti efektif "
        "dalam mengatasi permasalahan subjektivitas penilaian lapangan melalui penetapan tujuh kriteria "
        "sosial-ekonomi yang terstandarisasi. Sistem ini menghasilkan perankingan yang transparan dan "
        "akuntabel dengan mekanisme ambang batas skor kelayakan sebesar 0,50, sehingga setiap keputusan "
        "yang dihasilkan dapat dijelaskan secara matematis kepada seluruh pemangku kepentingan. Validasi "
        "melalui pengujian Black Box Testing mengonfirmasi bahwa seluruh fungsi aplikasi berjalan tanpa "
        "kegagalan maupun cacat kritikal, dan pengujian algoritma terhadap data sampel mencapai akurasi "
        "logika matematis 100%. Dengan demikian, sistem ini siap menjadi landasan transformasi digital "
        "proses seleksi bantuan sosial di Kecamatan Pondok Aren."
    )

    add_sub_heading(doc, "4.2 Saran")
    add_body(doc,
        "Guna memaksimalkan dampak sistem yang telah dibangun dan menjamin keberlanjutannya dalam jangka "
        "panjang, terdapat beberapa rekomendasi pengembangan yang perlu dipertimbangkan oleh pihak instansi "
        "maupun pengembang selanjutnya. Pertama, sistem sangat disarankan untuk diintegrasikan dengan "
        "aplikasi SIAK Terpusat agar validasi data kependudukan dapat dilakukan secara otomatis dan "
        "meminimalisasi potensi kesalahan input. Kedua, perlu dilakukan stress testing terhadap performa "
        "komputasi sistem menggunakan volume data berskala massal \u2014 yakni ribuan data warga \u2014 untuk "
        "memastikan stabilitasnya dalam kondisi operasional nyata. Ketiga, penerapan arsitektur multi-user "
        "dengan pemberian hak akses yang terdifferensiasi bagi staf di masing-masing kelurahan akan "
        "meningkatkan efisiensi dan keamanan pengelolaan data secara signifikan. Keempat, pengembangan "
        "modul dinamis yang memungkinkan penetapan standar kriteria berbeda untuk setiap variasi program "
        "bantuan \u2014 seperti BLT, PKH, dan BPNT \u2014 akan memperluas jangkauan kegunaan sistem. Kelima, "
        "migrasi basis data dari SQLite ke MySQL atau PostgreSQL sangat dianjurkan pada fase implementasi "
        "operasional penuh guna mencegah terjadinya database lock saat diakses secara bersamaan. Keenam "
        "dan terakhir, penyusunan user manual yang komprehensif perlu segera dilakukan untuk memastikan "
        "keberlanjutan operasional dan kemudahan pemeliharaan aplikasi oleh staf yang berwenang di masa "
        "mendatang."
    )

    # ========================================================================
    # SAVE
    # ========================================================================
    doc.save(OUTPUT_FILE)
    print(f"\n[OK] Laporan berhasil dibuat: {OUTPUT_FILE}")
    print(f"   Ukuran: {os.path.getsize(OUTPUT_FILE):,} bytes")


if __name__ == '__main__':
    generate_document()
